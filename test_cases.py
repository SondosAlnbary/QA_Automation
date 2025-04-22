import pytest
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains



# # in ex 6 - pip install python-docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from time import sleep

# @pytest.fixture
# yield driver
@pytest.fixture
def driver():
    driver = webdriver.Chrome()
    driver.maximize_window()
    # ----------------------------------------
    driver.get("https://www.terminalx.com/")
    sleep(5)
    yield driver
    driver.quit()

# -----------------------------------------------------------------------------------------
#     Test_Case -  1 : signup - create account
# ------------------------------------------------------------------------------------------

# def test_signup(driver):
#
#     driver.find_element(By.XPATH, "//button[@data-test-id = 'qa-header-login-button']").click()
#     sleep(3)
#
#     tab = driver.find_elements(By.CLASS_NAME,"tab_ZRBF")[1]
#     tab.click()
#     sleep(3)
#
#     f_name=driver.find_element(By.XPATH,"//input[@id = 'qa-register-firstname-input']")
#     f_name.send_keys("sara")
#     l_name=driver.find_element(By.XPATH,"//input[@id ='qa-register-lastname-input']")
#     l_name.send_keys("elreati")
#     email=driver.find_element(By.XPATH,"//input[@id ='qa-register-email-input']")
#     email.send_keys("autoyaso.97@gmail.com")
#     password=driver.find_element(By.XPATH,"//input[@id ='qa-register-password-input']")
#     password.send_keys("Test123$")
#     phone=driver.find_element(By.XPATH,"//input[@id ='qa-register-telephone-input']")
#     phone.send_keys("0534877201")
#     date_of_birth = driver.find_element(By.XPATH, "//input[@id ='qa-register-date_of_birth-input']")
#     date_of_birth.send_keys("12/03/1997")
#     sleep(5)
#     driver.find_element(By.XPATH, "//label[contains(text(),'נשים')]").click()
#     driver.find_element(By.XPATH, "//label[contains(text(),'גברים')]").click()
#     driver.find_element(By.XPATH, "//label[contains(text(),'ילדים')]").click()
#
#     sleep(1)
#     submit=driver.find_element(By.CLASS_NAME,"tx-link-a.dc-btn_1C6_.tx-link_29YD")
#     submit.click()
# -----------------------------------------------------------------------------------------
#     Test_Case -  2  : signin (login)
# ------------------------------------------------------------------------------------------
# def test_signin(driver):
#
#     driver.find_element(By.XPATH, "//button[@data-test-id = 'qa-header-login-button']").click()
#     sleep(2)
#     tab = driver.find_elements(By.CLASS_NAME, "tab_ZRBF")[0]
#     tab.click()
#     sleep(3)
#
#     email_input=driver.find_element(By.XPATH, "//input[@id = 'qa-login-email-input']")
#     email_input.send_keys("autoyaso.97@gmail.com")
#     pass_input=driver.find_element(By.XPATH, "//input[@ id = 'qa-login-password-input']")
#     pass_input.send_keys("Test123$")
#     signin_btn=driver.find_element(By.XPATH, "//button[@data-test-id = 'qa-login-submit']")
#     signin_btn.click()
#     sleep(5)

# ------------------------------------------------------------------------------------------
#     Test_Case -  3  :  search by Nike keyword
# ------------------------------------------------------------------------------------------

# def test_search(driver):
#
#     search_btn = driver.find_element(By.XPATH, "//button[@data-test-id = 'qa-header-search-button']")
#     search_btn.click()
#     search_input= driver.find_element(By.XPATH, "//input[@data-test='search-input']")
#
#     search_input.send_keys("NIKE")
#     search_input.send_keys(Keys.RETURN)
#
#     sleep(5)
#
#     images = driver.find_elements(By.CLASS_NAME, "image_3k9y")
#     assert len(images) > 0, "No products found on the page!"
#
#     i=1
#     for img in images:
#         src = img.get_attribute("src")
#         title = img.get_attribute("title")
#         alt = img.get_attribute("alt")
#         print(f"\n prodct - {i}")
#         print(f"Title: {title}")
#         print(f"Image: {src}")
#         print(f"Alt: {alt}")
#         i+=1
#         print("\n")
#         print("===" * 30)

# ------------------------------------------------------------------------------------------
#     Test_Case -  4  :  Get all brand names+logo images
# ------------------------------------------------------------------------------------------
#
def test_brands(driver):
    # Get all brand logo images
    brands = driver.find_elements(By.XPATH, '//div[contains(@class, "img-no-div-wrapper_VSJT")]//img')

    assert brands, "No brand logos found on the page"
    # Print brand names and image URLs
    i = 1
    for brand in brands:
        name = brand.get_attribute("alt")
        # link of brand img
        img = brand.get_attribute("src")
        # ASSERT: Brand name and image must be present
        assert name and name.strip(), f"Brand #{i} has missing or empty alt attribute"
        assert img and img.strip(), f"Brand #{i} has missing or empty image source"

        print(f"{i}. {name} -> {img}")
        i += 1

# ------------------------------------------------------------------------------------------
#     Test_Case -  5  :  Get discount_10 by email address
# ------------------------------------------------------------------------------------------

# def test_discount10(driver):
#     # Click on the 10% discount popup/button
#     discount_10 = driver.find_element(By.XPATH, "//div[@class='message_l1Lb']")
#     discount_10.click()
#     sleep(3)
#
#     # ASSERT: Email input should be visible
#     email_input = driver.find_element(By.XPATH, "//input[@name='email']")
#     assert email_input.is_displayed(), "Email input field is not visible"
#     email_input.send_keys("autoyaso.97@gmail.com")
#     sleep(2)
#
#     # Submit the form
#     submit_btn = driver.find_element(By.XPATH, "//button[@type = 'submit']")
#     assert submit_btn.is_enabled(), "Submit button is not clickable"
#     submit_btn.click()
#     sleep(3)
#
#     # ASSERT: Success or validation feedback appears
#     success_msgs = driver.find_elements(By.XPATH, "//div[contains(@class, 'form-feedback') or contains(text(), 'תודה') or contains(text(), 'Thank you')]")
#     assert success_msgs, "No confirmation message appeared after submitting email"

# ------------------------------------------------------------------------------------------
#     Test_Case -  6  : get all helpasking with answers and add in exe file
#  !!! notice : write in terminal =>  - pip install python-docx, and libs in frist page to import doc
# ------------------------------------------------------------------------------------------

# def test_help_asking(driver):
#
#
#     driver.get("https://www.terminalx.com/help")
#     sleep(3)
#
#     # Find all FAQ items
#     faq_items = driver.find_elements(By.XPATH, "//ul[@class='list_2swX']/li")
#
#     # Create Word document
#     doc = Document()
#     doc.add_heading("TerminalX - שאלות ותשובות", 0)
#
#     for i in range(len(faq_items)):
#         faq_items = driver.find_elements(By.XPATH, "//ul[@class='list_2swX']/li")
#         question_el = faq_items[i].find_element(By.XPATH, ".//div[contains(@class,'question')]")
#         question_el.click()
#         sleep(0.5)
#
#         question_text = question_el.text.strip()
#         answer_paragraphs = faq_items[i].find_elements(By.XPATH, ".//div[contains(@class,'answer')]//p")
#         full_answer = " ".join([p.text.strip() for p in answer_paragraphs if p.text.strip()])
#
#         # Add question
#         q = doc.add_paragraph()
#         q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#         run_q = q.add_run(f"{question_text}")
#         run_q.bold = True
#         run_q.font.size = Pt(12)
#
#         # Add answer
#         a = doc.add_paragraph()
#         a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#         run_a = a.add_run(f"{full_answer}")
#         run_a.font.size = Pt(11)
#
#         doc.add_paragraph("")  # Spacer
#
#     # Save document
#     doc.save("TerminalX_help_asking.docx")

# ------------------------------------------------------------------------------------------
#     Test_Case -  7  :info by all BEST CATEGORIES in women section (alt+linkimg)
# ------------------------------------------------------------------------------------------
# def test_cate_women(driver):
#
#     # Grab BEST CATEGORIES cards
#     cards = driver.find_elements(By.XPATH, "//div[@class='img-no-div-wrapper_VSJT']")
#
#     # ASSERT: Cards must exist
#     assert cards, "No BEST CATEGORIES cards found in WOMEN section"
#
#     categories = []
#
#     for card in cards:
#         try:
#             img = card.find_element(By.TAG_NAME, "img")
#             alt = img.get_attribute("alt")
#             src = img.get_attribute("src")
#             categories.append({"alt": alt, "img": src})
#         except Exception as e:
#             print(f"Error extracting category info: {e}")
#             continue
#
#     # ASSERT: At least one category extracted
#     assert categories, "No categories extracted from WOMEN section"
#
#     # Output the BEST CATEGORIES information
#     print(f"\n=== WOMEN BEST CATEGORIES ===")
#     for c in categories:
#         print(f"Category: {c['alt']}")
#         print(f"Image: {c['img']}")
#         print("-" * 40)
    #
# ------------------------------------------------------------------------------------------
#     Test_Case -  8  :info by all BEST CATEGORIES in men section (alt+linkimg)
# ------------------------------------------------------------------------------------------
# def test_cate_men(driver):
#     # Find the nav button for the "MEN" section
#     men_link = driver.find_element(By.XPATH, "//a[@href='men']")
#     men_link.click()
#     sleep(5)
#
#     # ASSERT: Verify navigation to MEN section
#     assert "/men" in driver.current_url, "Failed to navigate to MEN section"
#
#     # Grab BEST CATEGORIES
#     cards = driver.find_elements(By.XPATH, "//div[@class='img-no-div-wrapper_VSJT']")
#
#     # ASSERT: Ensure cards exist
#     assert cards, "No BEST CATEGORIES cards found on MEN page"
#
#     categories = []
#
#     for card in cards:
#         try:
#             img = card.find_element(By.TAG_NAME, "img")
#             alt = img.get_attribute("alt")
#             src = img.get_attribute("src")
#             categories.append({"alt": alt, "img": src})
#         except Exception as e:
#             print(f"Error reading category card: {e}")
#             continue
#
#     # ASSERT: Ensure at least one category was extracted
#     assert categories, "No categories extracted from BEST CATEGORIES section"
#
#     # Output the BEST CATEGORIES information
#     print(f"\n=== MEN BEST CATEGORIES ===")
#     for c in categories:
#         print(f"Category: {c['alt']}")
#         print(f"Image: {c['img']}")
#         print("-" * 40)

# ------------------------------------------------------------------------------------------
#     Test_Case -  9  :get info by all BEST CATEGORIES  (alt+linkimg)  in sections for
#                                WOMEN,MEN,KIDS,LIVING
# ------------------------------------------------------------------------------------------

# def test_best_cat1(driver):
#     # Define category hrefs mapped to readable names
#     category_hrefs = {
#         "WOMEN": "women",
#         "MEN": "men",
#         "KIDS": "kids",
#         "LIVING": "/home-lifestyle"
#     }
#
#     # Loop through each category
#     for label, href in category_hrefs.items():
#         print(f"\n section - {label}")
#
#         xpath = f"//nav[contains(@class,'universeSelectorWrapper_213P')]/a[contains(@href, '{href}')]"
#         category_link = driver.find_element(By.XPATH, xpath)
#         category_link.click()
#         sleep(3)
#
#         # Find all image elements
#         images = driver.find_elements(By.TAG_NAME, "img")
#
#         valid_img = 0
#         for img in images:
#             alt = img.get_attribute("alt")
#             src = img.get_attribute("src")
#             if alt and src:
#                 print(f"name_cat: {alt}\nimg: {src}\n")
#                 valid_img += 1
#
#         # Assert at least 1 valid image was found
#         assert valid_img > 0, f"No valid images found in category: {label}"
#         print(f"Found {valid_img} valid category in {label}")

# ------------------------------------------------------------------------------------------
#     Test_Case -  10  :get info by all BEST CATEGORIES  (alt+linkimg)  in sections for
# JUST LANDED,ON SALE, נשים,גברים, ילדים, בייבי, מותגים, SPORTS, BEAUTY, HOME, JEWELRY,WELLNESS
# ------------------------------------------------------------------------------------------

# def test_best_cats2(driver):
#     # Category list
#     categories = [
#         "JUST LANDED", "ON SALE", "נשים", "גברים", "ילדים", "בייבי",
#         "מותגים", "SPORTS", "BEAUTY", "HOME", "JEWELRY", "WELLNESS"
#     ]
#
#     # Loop through each category
#     for category in categories:
#         try:
#             print(f"\n\ncategory: {category}")
#
#             # Find and click the category button
#             category_btn = driver.find_element(By.XPATH, f"//a[contains(text(), '{category}')]")
#             ActionChains(driver).move_to_element(category_btn).click().perform()
#             sleep(4)
#
#             images = driver.find_elements(By.TAG_NAME, "img")
#
#             for img in images:
#                 alt = img.get_attribute("alt")
#                 src = img.get_attribute("src")
#                 if alt and src:
#                     print(f"category - {category}")
#                     print(f"name: {alt} \nsrc: {src}\n")
#
#         except Exception as e:
#             print(f"Error in category '{category}': {e}")

# ------------------------------------------------------------------------------------------
#     Test_Case -  11  : sort product by מבצע option
# ------------------------------------------------------------------------------------------

# def test_sort_just_landed_by_sale(driver):
#     #Scroll to the JUST LANDED section and click on it
#     just_landed = driver.find_element(By.XPATH, "//a[@href='/justlanded']")
#     assert just_landed is not None, "'JUST LANDED' link is not found."
#     just_landed.click()
#     sleep(4)
#     #
#     # Verify the "JUST LANDED" page has loaded
#     assert "JUST LANDED" in driver.title, "'JUST LANDED' page did not load correctly."
#     #
#     # Click on the sort dropdown (3 dots icon or select box)
#     sort_icon = driver.find_element(By.XPATH, "//div[@class='select-white_3SVZ']")
#     #
#     sort_icon.click()
#     sleep(4)
#
#     sort_dropdown = driver.find_element(By.XPATH, "//select[@name='sortField']//option[text()='מבצע']")
#     sort_dropdown.click()
#
#     sleep(4)
#
#     assert "מבצע" in driver.page_source, "'מבצע' sort option was not applied successfully."
#     print(" Sorted 'JUST LANDED' products by 'מבצע' (Sale) successfully.")



# ------------------------------------------------------------------------------------------
#     Test_Case -  12  sort by Price: Low to High" option in on-sale page
# ------------------------------------------------------------------------------------------
# def test_pyl_h(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     sort_btn = driver.find_element(By.XPATH, "//select[@name='sortField']")
#     sort_btn.click()
#     sleep(1)
#
#     # Click on Price: Low to High" option
#     price_asc_option =driver.find_element(By.XPATH, "//select[@name='sortField']" and "//option[@value='price_asc']")
#     price_asc_option.click()
#     sleep(5)

# ------------------------------------------------------------------------------------------
#     Test_Case -  13  sort by Price:  High To Low " option in on-sale page
# ------------------------------------------------------------------------------------------

# def test_pyh_l(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     sort_btn = driver.find_element(By.XPATH, "//select[@name='sortField']")
#     sort_btn.click()
#     sleep(1)
#
#     # Click on Price:  High to Low" option
#     price_desc_option =driver.find_element(By.XPATH, "//select[@name='sortField']" and "//option[@value='price_desc']")
#     price_desc_option.click()
#     sleep(5)

# ------------------------------------------------------------------------------------------
#     Test_Case -  14  sort by :  BEST OFFER  option in on-sale page
# ------------------------------------------------------------------------------------------

# def test_best_offer(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     sort_btn = driver.find_element(By.XPATH, "//select[@name='sortField']")
#     sort_btn.click()
#     sleep(1)
#
#
#     price_best_offer =driver.find_element(By.XPATH, "//select[@name='sortField']" and "//option[@value='bestsellers']")
#     price_best_offer.click()
#     sleep(5)

# ------------------------------------------------------------------------------------------
#     Test_Case -  15  print info for 10 products in on-sale section
# ------------------------------------------------------------------------------------------

# def test_data_pros(driver):
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     # Find product cards
#     products = driver.find_elements(By.XPATH, "//li[contains(@class, 'listing-product_3mjp')]")
#
#     # Loop through a few products
#     for product in products[:10]:
#         try:
#             # Final price
#             price = product.find_element(By.CLASS_NAME, "final-price_8CiX").text.strip()
#
#             # Brand
#             brand = product.find_element(By.XPATH, ".//div[contains(@class, 'right_1o65')]/span").text.strip()
#
#             # Product name
#             title_elem = product.find_element(By.XPATH, ".//a[contains(@class, 'title_3ZxJ')]")
#             title = title_elem.text.strip()
#
#             # Link (relative href)
#             link = title_elem.get_attribute("href")
#
#             # Color swatch title (tooltip)
#             color = product.find_element(By.XPATH, ".//div[contains(@class,'color-item_1Y2Y')]").get_attribute("title")
#
#             print(f"{brand} - {title}")
#             print(f"Price - {price}")
#             print(f"olor - {color}")
#             print(f"Link - {link}")
#             print("----")
#         except Exception as e:
#             print("Skipped one product due to missing info.")
#             continue

# ------------------------------------------------------------------------------------------
#     Test_Case -  16  print cheapest price for product on-sale section
# ------------------------------------------------------------------------------------------

# def test_cheap_price(driver):
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     sort_btn = driver.find_element(By.XPATH, "//select[@name='sortField']")
#     sort_btn.click()
#     sleep(1)
#
#     # Click on Price: Low to High" option
#     price_asc_option =driver.find_element(By.XPATH, "//select[@name='sortField']" and "//option[@value='price_asc']")
#     price_asc_option.click()
#     sleep(5)
#
#     # Scroll to load more products
#     driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
#     sleep(10)
#
#     # Find product cards
#     products = driver.find_elements(By.XPATH, "//li[contains(@class, 'listing-product_3mjp')]")
#
#     # Track cheapest product
#     min_price = float('inf')
#     cheapest_product = {}
#
#     # Loop through products
#     for product in products:
#         try:
#             # Price
#             price_text = product.find_element(By.CLASS_NAME, "final-price_8CiX").text.strip()
#             price_value = float(price_text.replace("₪", "").replace(",", "").strip())
#
#             if price_value < min_price:
#                 # Brand
#                 brand = product.find_element(By.XPATH, ".//div[contains(@class, 'right_1o65')]/span").text.strip()
#
#                 # Title
#                 title_elem = product.find_element(By.XPATH, ".//a[contains(@class, 'title_3ZxJ')]")
#                 title = title_elem.text.strip()
#
#                 # Link
#                 link = title_elem.get_attribute("href")
#
#                 # Color
#                 color = product.find_element(By.XPATH, ".//div[contains(@class,'color-item_1Y2Y')]").get_attribute(
#                     "title")
#
#                 # Save cheapest product
#                 min_price = price_value
#                 cheapest_product = {
#                     "price": price_text,
#                     "brand": brand,
#                     "title": title,
#                     "link": link,
#                     "color": color
#                 }
#         except:
#             continue
#
#     # Print result
#     if cheapest_product:
#         print("\nCheapest Product on Sale:")
#         print(f"{cheapest_product['brand']} - {cheapest_product['title']}")
#         print(f"Price: {cheapest_product['price']}")
#         print(f"Color: {cheapest_product['color']}")
#         print(f"Link: {cheapest_product['link']}")
#     else:
#         print("No valid products found.")
# ------------------------------------------------------------------------------------------
#     Test_Case -  17  print expensive price for product on-sale section
# ------------------------------------------------------------------------------------------

# def test_expensive_price(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#     sort_btn = driver.find_element(By.XPATH, "//select[@name='sortField']")
#     sort_btn.click()
#     sleep(1)
#
#     # Click on Price:  High to Low" option
#     price_desc_option =driver.find_element(By.XPATH, "//select[@name='sortField']" and "//option[@value='price_desc']")
#     price_desc_option.click()
#     sleep(5)
#     # Scroll to load more products
#     driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
#     sleep(15)
#     # Find product cards
#     products = driver.find_elements(By.XPATH, "//li[contains(@class, 'listing-product_3mjp')]")
#
#     # Track most expensive
#     max_price = 0
#     expensive_product = {}
#
#     # Loop through products
#     for product in products:
#         try:
#             # Price
#             price_text = product.find_element(By.CLASS_NAME, "final-price_8CiX").text.strip()
#             price_value = float(price_text.replace("₪", "").replace(",", "").strip())
#
#             if price_value > max_price:
#                 # Brand
#                 brand = product.find_element(By.XPATH,
#                                              ".//div[contains(@class, 'right_1o65')]/span").text.strip()
#
#                 # Title
#                 title_elem = product.find_element(By.XPATH, ".//a[contains(@class, 'title_3ZxJ')]")
#                 title = title_elem.text.strip()
#
#                 # Link
#                 link = title_elem.get_attribute("href")
#
#                 # Color
#                 color = product.find_element(By.XPATH,
#                                              ".//div[contains(@class,'color-item_1Y2Y')]").get_attribute(
#                     "title")
#
#                 # Save most expensive
#                 max_price = price_value
#                 expensive_product = {
#                     "price": price_text,
#                     "brand": brand,
#                     "title": title,
#                     "link": link,
#                     "color": color
#                 }
#         except:
#             continue
#
#     # Print result
#     if expensive_product:
#         print("\nMost Expensive Product on Sale:")
#         print(f" {expensive_product['brand']} - {expensive_product['title']}")
#         print(f"Price: {expensive_product['price']}")
#         print(f"Color: {expensive_product['color']}")
#         print(f"Link: {expensive_product['link']}")
#     else:
#         print("No valid products found.")

# ------------------------------------------------------------------------------------------
#     Test_Case -  18  add for mylist products on-sale section
# ------------------------------------------------------------------------------------------

# def test_my_list(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(3)
#
#     # Scroll to ensure all quick-views load
#     driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
#     sleep(4)
#
#     # Find the wishlist "REMOVE"/heart icon containers (not button directly)
#     heart_containers = driver.find_elements(
#         By.XPATH,
#         "//div[contains(@class, 'quick-view_2qdz')]//div[contains(@class, 'btn-my_list_2EOz')]"
#     )
#
#     print(f"Found {len(heart_containers)} heart icon containers")
#
#     # Click the button inside each container
#     for idx, container in enumerate(heart_containers):
#         try:
#             btn = container.find_element(By.TAG_NAME, "button")
#             driver.execute_script("arguments[0].click();", btn)
#             print(f"Clicked heart (REMOVE) button #{idx + 1}")
#             sleep(1)
#         except Exception as e:
#             print(f"Could not click heart #{idx + 1}: {e}")
# #
# ------------------------------------------------------------------------------------------
#     Test_Case -  19  : make quick view for products
# ------------------------------------------------------------------------------------------

# def test_quick_view(driver):
#         # Navigate to On Sale
#         on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#         on_sale_btn.click()
#         sleep(3)
#
#         quick_view_buttons = driver.find_elements(
#             By.XPATH,
#             "//div[contains(@class, 'btn-quick-view_2SXw')]/button"
#         )
#
#         print(f"Found {len(quick_view_buttons)} Quick View buttons")
#
#         for idx, btn in enumerate(quick_view_buttons):
#             try:
#                 driver.execute_script("arguments[0].scrollIntoView(true);", btn)
#                 sleep(1)
#                 driver.execute_script("arguments[0].click();", btn)
#                 print(f"Opened Quick View #{idx + 1}")
#                 sleep(2)
#
#                 # Simulate hover on zoom layer
#                 try:
#                     zoom_layer = driver.find_element(
#                         By.XPATH,
#                         "//div[contains(@class, 'zoom-layer_1azI')]"
#                     )
#                     actions = ActionChains(driver)
#                     actions.move_to_element(zoom_layer).perform()
#                     print(f"Hovered on zoom layer #{idx + 1}")
#                     sleep(1.5)
#                 except Exception as zoom_err:
#                     print(f"Could not hover on zoom layer #{idx + 1}: {zoom_err}")
#
#                 # Close modal
#                 close_btn = driver.find_element(
#                     By.XPATH,
#                     "//button[contains(@class, 'close') or contains(@aria-label, 'Close')]"
#                 )
#                 driver.execute_script("arguments[0].click();", close_btn)
#                 print(f"Closed Quick View #{idx + 1}")
#                 sleep(1)
#
#             except Exception as e:
#                 print(f"Error at Quick View #{idx + 1}: {e}")

# ------------------------------------------------------------------------------------------
#     Test_Case -  20  : add products to card in on-sale section
# ------------------------------------------------------------------------------------------
#
# def test_add_cart(driver):
#
#     # Hover over WOMEN
#     women_menu = driver.find_element(By.XPATH, "//a[contains(@href, '/women')]")
#     women_menu.click()
#     sleep(2)
#
#     # Click on T-Shirts
#     shirts_link = driver.find_element(By.XPATH, "//img[@alt='SHIRTS']")
#     shirts_link.click()
#     sleep(5)
#
#     # Find first 5 product name links
#     product_links = driver.find_elements(By.XPATH, "//a[@data-test-id='qa-product-link']")
#     assert len(product_links) >= 5, "Less than 5 products found"
#
#     # Get hrefs to avoid stale elements after navigation
#     product_urls = [link.get_attribute("href") for link in product_links[:5]]
#
#     # Loop through the first 5 product URLs
#     for i in range(5):
#         driver.get(product_urls[i])
#         sleep(4)
#
#         # Find and click the first available size (non-disabled)
#         size_elements = driver.find_elements(By.XPATH,"//div[@data-test-id='qa-size-item']"
#         )
#         assert size_elements, f"No available size found on product {i + 1}"
#
#         # Click the first available size
#         size_elements[0].click()
#         sleep(2)
#
#         for size in sizes:
#             for size_element in size_elements:
#                 if size_element.text.strip() == size:  # Check if the size matches
#                     print(f"Size {size} available, clicking it for product {i + 1}.")
#                     size_element.click()
#                     sleep(2)  # Ensure the click is registered
#
#
#         # Click Add to Cart
#         add_to_cart_btn = driver.find_element(By.XPATH, "//button[@data-test-id='qa-add-to-cart-button']")
#         add_to_cart_btn.click()
#         sleep(3)
#
#         print(f"Product {i + 1} added to cart.")
#
#         # Verify cart count
#     cart_badge = driver.find_element(By.XPATH, "//span[contains(@class,'counter-number')]")
#     cart_count = int(cart_badge.text)
#     assert cart_count >= 5, f"Expected 5 items in cart, found {cart_count}"

# ------------------------------------------------------------------------------------------
#     Test_Case -  21  : total product in women -> T-shirt section
# ------------------------------------------------------------------------------------------

# def test_total_items(driver):
#     # Hover over WOMEN
#     women_menu = driver.find_element(By.XPATH, "//a[contains(@href, '/women')]")
#     women_menu.click()
#     sleep(2)
#
#     # Click on T-Shirts
#     shirts_link = driver.find_element(By.XPATH, "//img[@alt='SHIRTS']")
#     shirts_link.click()
#     sleep(5)
#
#     items_num = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#     items_text = items_num.text.strip()
#     print(f"\n\nT-shirt - {items_text} items")

# ------------------------------------------------------------------------------------------
#     Test_Case -  22  : total product in JUST LANDED section in sidebar
# ------------------------------------------------------------------------------------------
# def test_total_items1(driver):
#
#     just_landed_btn = driver.find_element(By.XPATH, "//a[@href='/justlanded']")
#     just_landed_btn.click()
#     sleep(2)
#
#     # Fixed category list
#     categories = ["BEAUTY", "HOME", "TEEN", "WELLNESS", "בייבי", "גברים", "ילדים", "נשים"]
#
#     category_data = []
#     loop_count = 0
#
#     for cat in categories:
#         loop_count += 1
#         print(f"section - '{cat}'")
#
#         cat_element = driver.find_element(By.XPATH, f"//ol/li/a[text()='{cat}']")
#         assert cat_element is not None, f"Element not found for '{cat}'"
#
#         cat_element.click()
#         sleep(4)
#
#         # Get item count
#         total_element = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#         total_text = total_element.text.strip()
#         assert total_text != "", f"No total item count for '{cat}'"
#
#         category_data.append((cat, total_text))
#         print(f"Tot_item - {total_text}")

# ------------------------------------------------------------------------------------------
#     Test_Case -  23  : total product in on-sale section - with range price 500-800
# ------------------------------------------------------------------------------------------

# def test_price_range500_800(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(2)
#
#     min_slider = driver.find_element(By.XPATH, "//span[@data-index='0']//input[@type='range']")
#     max_slider = driver.find_element(By.XPATH, "//span[@data-index='1']//input[@type='range']")
#
#     driver.execute_script("arguments[0].scrollIntoView(true);", min_slider)
#     sleep(1)
#
#     driver.execute_script("""
#             arguments[0].value = 500;
#             arguments[0].setAttribute('aria-valuenow', '500');
#             arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
#             arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
#         """, min_slider)
#
#
#     driver.execute_script("""
#             arguments[0].value = 800;
#             arguments[0].setAttribute('aria-valuenow', '800');
#             arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
#             arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
#         """, max_slider)
#
#     sleep(10)
#
#     total_element = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#     total_text = total_element.text
#     print(f"\n\nTotal products in range 500-800: {total_text}")
#     assert total_text != "", "Total products text is empty — filter may not have applied correctly."
#     assert any(char.isdigit() for char in total_text), f"No product count found in: {total_text}"
# ------------------------------------------------------------------------------------------
#     Test_Case -  24  : total product in on-sale section - with one-size
# ------------------------------------------------------------------------------------------

# def test_pro_on_size(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(2)
#
#     size=driver.find_element(By.XPATH, "//div[@class='header_OqIf']/h4[contains(text(), 'מידה')]")
#     size.click()
#     sleep(2)
#
#     # Click the "OneSize" filter link
#     one_size = driver.find_element(By.XPATH,"//ol/li[@class='filter-item_wzYv'][13]")
#     print(f"\nfilter by option - {one_size.text}")
#     one_size.click()
#     sleep(5)
#
#     total_element = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#     total_text = total_element.text
#     print(f"\n\ntotal product - {total_text}")
#
#     assert total_text != "", "Total products text is empty — filter may not have applied correctly."
#     assert any(char.isdigit() for char in total_text), f"No product count found in: {total_text}"

# ------------------------------------------------------------------------------------------
#     Test_Case -  25  : total product in on-sale section - with 50%
# ------------------------------------------------------------------------------------------

# def test_pro_on_size(driver):
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(2)
#
#     sale = driver.find_element(By.XPATH, "//div[@class='header_OqIf']/h4[contains(text(), 'מבצע')]")
#     sale.click()
#     sleep(2)
#
#     # Click the "OneSize" filter link
#     percent_50 = driver.find_element(By.XPATH, "//ol/li[@class='filter-item_wzYv'][18]")
#     print(f"\nfilter by option - {percent_50.text}")
#     percent_50.click()
#     sleep(5)
#
#     total_element = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#     total_text = total_element.text
#     print(f"\n\ntotal product - {total_text}")
#
#     assert total_text != "", "Total products text is empty — filter may not have applied correctly."
#     assert any(char.isdigit() for char in total_text), f"No product count found in: {total_text}"

# ------------------------------------------------------------------------------------------
#     Test_Case -  26  : total product in brand section - with FILA option
# ------------------------------------------------------------------------------------------
# def test_brand(driver):
#
#     on_sale_btn = driver.find_element(By.XPATH, "//a[@href='/on-sale']")
#     on_sale_btn.click()
#     sleep(2)
#
#     brand_btn = driver.find_element(By.XPATH, "//div[@class='header_OqIf']/h4[contains(text(),'מותג')]")
#     brand_btn.click()
#     sleep(2)
#
#     # Click the "OneSize" filter link
#     brand = driver.find_element(By.XPATH, "//li[@class='filter-item_wzYv brand_2f3C']/a[text()='FILA']")
#     print(f"\n\nbrand - {brand.text}")
#     brand.click()
#     sleep(5)
#
#     total_element = driver.find_element(By.XPATH, "//span[@data-test='search-totals']")
#     total_text = total_element.text
#     print(f"\n\ntotal product - {total_text}")
#
#     assert total_text != "", "Total products text is empty — filter may not have applied correctly."
#     assert any(char.isdigit() for char in total_text), f"No product count found in: {total_text}"



















